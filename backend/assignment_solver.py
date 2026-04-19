import os
import re
import winreg
from groq import Groq
from dotenv import load_dotenv

load_dotenv()

# ─────────────────────────────────────────
# Groq client (reuses existing env key)
# ─────────────────────────────────────────
_client = Groq(api_key=os.getenv("GROQ_API_KEY"))

# ─────────────────────────────────────────
# Command trigger phrases
# ─────────────────────────────────────────
ASSIGNMENT_TRIGGERS = [
    "complete assignment",
    "complete my assignment",
    "do my assignment",
    "do assignment",
    "solve assignment",
    "solve my assignment",
    "finish assignment",
    "finish my assignment",
    "answer assignment",
]


# ─────────────────────────────────────────
# Get REAL Desktop path (handles OneDrive)
# ─────────────────────────────────────────

def _get_real_desktop() -> str:
    """
    Uses Windows Registry to get the actual Desktop path.
    This correctly resolves OneDrive-redirected Desktops like:
      C:\\Users\\User\\OneDrive\\Desktop
    instead of the wrong:
      C:\\Users\\User\\Desktop
    """
    try:
        key = winreg.OpenKey(
            winreg.HKEY_CURRENT_USER,
            r"Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders"
        )
        desktop, _ = winreg.QueryValueEx(key, "Desktop")
        winreg.CloseKey(key)
        if os.path.isdir(desktop):
            return desktop
    except Exception:
        pass

    # Fallback: try both standard and OneDrive paths
    home = os.path.expanduser("~")
    candidates = [
        os.path.join(home, "OneDrive", "Desktop"),
        os.path.join(home, "Desktop"),
    ]
    for path in candidates:
        if os.path.isdir(path):
            return path

    return os.path.join(home, "Desktop")  # last resort


# ─────────────────────────────────────────
# Common search root folders
# ─────────────────────────────────────────

def _get_search_roots() -> list[str]:
    """
    Returns all meaningful locations to search for a file:
    Desktop (real), Documents, Downloads, OneDrive root,
    and all available drive roots.
    """
    home = os.path.expanduser("~")
    roots = []

    # High-priority common folders
    priority_dirs = [
        _get_real_desktop(),
        os.path.join(home, "Documents"),
        os.path.join(home, "Downloads"),
        os.path.join(home, "OneDrive"),
        os.path.join(home, "OneDrive - Personal"),
        os.path.join(home, "Desktop"),
        home,
    ]
    for d in priority_dirs:
        if os.path.isdir(d) and d not in roots:
            roots.append(d)

    # All available drive roots (C:\, D:\, etc.)
    import string
    for letter in string.ascii_uppercase:
        drive = f"{letter}:\\"
        if os.path.isdir(drive) and drive not in roots:
            roots.append(drive)

    return roots


def is_assignment_command(command: str) -> bool:
    """Return True if the command is an assignment completion request."""
    cmd = command.lower().strip()
    return any(trigger in cmd for trigger in ASSIGNMENT_TRIGGERS)


def parse_filename(command: str) -> str | None:
    """
    Extract the filename from the command.

    Examples:
      "complete assignment homework.txt"        → "homework.txt"
      "do my assignment called math_task.pdf"   → "math_task.pdf"
      "solve assignment file notes.docx"        → "notes.docx"
      "complete my assignment science.pdf"      → "science.pdf"
    """
    cmd = command.strip()

    # Remove trigger phrase (longest-first to avoid partial matches)
    for trigger in sorted(ASSIGNMENT_TRIGGERS, key=len, reverse=True):
        if trigger in cmd.lower():
            idx = cmd.lower().find(trigger)
            cmd = cmd[idx + len(trigger):].strip()
            break

    # Strip filler words
    filler = r"^(called|named|file|for|my|the|:)\s+"
    cmd = re.sub(filler, "", cmd, flags=re.IGNORECASE).strip()

    return cmd if cmd else None


# ─────────────────────────────────────────
# File finder — searches Desktop + entire filesystem
# ─────────────────────────────────────────

def _match_filename(filename: str, candidate: str) -> bool:
    """
    Returns True if `candidate` matches `filename`.
    Supports:
      - Exact match
      - Case-insensitive match
      - Match without extension (e.g. "Experiment_8" matches "Experiment_8.docx")
    """
    lower_fn   = filename.lower()
    lower_cand = candidate.lower()

    if lower_cand == lower_fn:
        return True

    # Match by base name (no extension)
    base_fn   = os.path.splitext(lower_fn)[0]
    base_cand = os.path.splitext(lower_cand)[0]
    if base_cand == base_fn and not os.path.splitext(lower_fn)[1]:
        # User gave no extension → match by base name only
        return True

    return False


def _find_file(filename: str) -> str | None:
    """
    Searches for `filename` across:
    1. The real Desktop (Registry-resolved, handles OneDrive)
    2. Common user folders: Documents, Downloads, OneDrive
    3. All drive roots (recursive walk, skips system dirs)

    Returns the full path of the first match, or None.
    """
    lower_fn  = filename.lower()
    base_fn   = os.path.splitext(lower_fn)[0]
    has_ext   = bool(os.path.splitext(lower_fn)[1])

    # Skip these directories to avoid wasting time
    SKIP_DIRS = {
        "windows", "system32", "syswow64", "program files",
        "program files (x86)", "programdata", "appdata",
        "$recycle.bin", "recovery", "perflogs",
        "node_modules", ".git", "__pycache__",
    }

    roots = _get_search_roots()
    searched_paths = set()

    for root in roots:
        try:
            for dirpath, dirnames, filenames in os.walk(root):
                # Avoid duplicate walks
                real_dirpath = os.path.realpath(dirpath)
                if real_dirpath in searched_paths:
                    dirnames.clear()
                    continue
                searched_paths.add(real_dirpath)

                # Prune skip dirs in-place to prevent descending
                dirnames[:] = [
                    d for d in dirnames
                    if d.lower() not in SKIP_DIRS
                ]

                for fname in filenames:
                    lower_fname = fname.lower()
                    base_fname  = os.path.splitext(lower_fname)[0]

                    # Exact case-insensitive match
                    if lower_fname == lower_fn:
                        return os.path.join(dirpath, fname)

                    # Match by base name when user didn't give extension
                    if not has_ext and base_fname == base_fn:
                        return os.path.join(dirpath, fname)

        except PermissionError:
            continue
        except Exception:
            continue

    return None


# ─────────────────────────────────────────
# File reader — supports multiple formats
# ─────────────────────────────────────────

def _read_file(filepath: str) -> str:
    """
    Read content from a file. Supports:
      - Plain text: .txt .md .py .js .ts .html .css .json .csv .xml
      - Word docs:  .docx  (python-docx)
      - PDFs:       .pdf   (PyMuPDF)
      - Excel:      .xlsx  (openpyxl)
    """
    ext = os.path.splitext(filepath)[1].lower()

    # ── Plain text ────────────────────────────────────────────────────────────
    plain_text_exts = {
        ".txt", ".md", ".py", ".js", ".ts", ".jsx", ".tsx",
        ".html", ".css", ".json", ".csv", ".xml", ".yaml", ".yml",
        ".java", ".c", ".cpp", ".cs", ".go", ".rs", ".rb", ".php",
    }
    if ext in plain_text_exts:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    # ── Word (.docx) ──────────────────────────────────────────────────────────
    if ext == ".docx":
        try:
            import docx
            doc = docx.Document(filepath)
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            raise RuntimeError("Run: pip install python-docx")

    # ── PDF ───────────────────────────────────────────────────────────────────
    if ext == ".pdf":
        try:
            import fitz
            pdf = fitz.open(filepath)
            text = "".join(page.get_text() for page in pdf)
            pdf.close()
            return text
        except ImportError:
            raise RuntimeError("Run: pip install PyMuPDF")

    # ── Excel (.xlsx / .xls) ─────────────────────────────────────────────────
    if ext in (".xlsx", ".xls"):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(filepath, data_only=True)
            lines = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                lines.append(f"=== Sheet: {sheet} ===")
                for row in ws.iter_rows(values_only=True):
                    lines.append(
                        "\t".join(str(c) if c is not None else "" for c in row)
                    )
            return "\n".join(lines)
        except ImportError:
            raise RuntimeError("Run: pip install openpyxl")

    raise RuntimeError(
        f"Unsupported file type '{ext}'. "
        "Supported: .txt .md .py .js .html .pdf .docx .xlsx .csv .json"
    )


# ─────────────────────────────────────────
# AI solver
# ─────────────────────────────────────────

def _ask_ai(assignment_content: str, filename: str) -> str:
    """Send assignment content to Groq and return the completed answer."""
    system_prompt = (
        "You are an expert academic assistant. "
        "The user will give you the content of their assignment file. "
        "Your job is to COMPLETE the assignment thoroughly and correctly. "
        "Provide well-structured, detailed, and accurate answers. "
        "If it is a set of questions, answer every question clearly. "
        "If it is an essay or report prompt, write a full essay or report. "
        "If it is a coding assignment, write the complete working code. "
        "Format your response professionally."
    )

    user_prompt = (
        f"Here is the content of my assignment file '{filename}':\n\n"
        f"{'='*60}\n"
        f"{assignment_content}\n"
        f"{'='*60}\n\n"
        "Please complete this assignment fully and correctly."
    )

    response = _client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_prompt},
        ],
        temperature=0.5,
        max_tokens=8000,
    )
    return response.choices[0].message.content.strip()


def _save_result(source_filepath: str, filename: str, result: str) -> str:
    """
    Save completed assignment as a formatted .docx Word document.
    - Saved next to the original file
    - A copy also placed on the real Desktop for easy access
    """
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    base     = os.path.splitext(filename)[0]
    out_name = f"{base}_completed.docx"

    source_dir = os.path.dirname(source_filepath)
    out_path   = os.path.join(source_dir, out_name)

    def _build_doc() -> docx.Document:
        doc = docx.Document()

        # ── Title ─────────────────────────────────────────────────────────────
        title = doc.add_heading(f"Completed Assignment", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)   # dark blue

        # ── Source file subtitle ───────────────────────────────────────────────
        sub = doc.add_paragraph(f"Source file: {filename}")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in sub.runs:
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

        doc.add_paragraph("")   # spacer

        # ── Body — split on double newlines into paragraphs ───────────────────
        sections = result.split("\n\n")
        for section in sections:
            section = section.strip()
            if not section:
                continue

            lines = section.split("\n")
            first_line = lines[0].strip()

            # Detect markdown-style headings (## or **bold title**)
            if first_line.startswith("##"):
                doc.add_heading(first_line.lstrip("#").strip(), level=2)
                body = "\n".join(lines[1:]).strip()
                if body:
                    doc.add_paragraph(body)

            elif first_line.startswith("#"):
                doc.add_heading(first_line.lstrip("#").strip(), level=1)
                body = "\n".join(lines[1:]).strip()
                if body:
                    doc.add_paragraph(body)

            elif first_line.startswith("**") and first_line.endswith("**"):
                # Bold heading line
                p = doc.add_paragraph()
                run = p.add_run(first_line.strip("*"))
                run.bold = True
                run.font.size = Pt(12)
                body = "\n".join(lines[1:]).strip()
                if body:
                    doc.add_paragraph(body)

            else:
                doc.add_paragraph(section)

        return doc

    # Save next to the original file
    doc = _build_doc()
    doc.save(out_path)

    # Also copy to real Desktop for convenience
    desktop   = _get_real_desktop()
    desk_path = os.path.join(desktop, out_name)
    if os.path.realpath(source_dir) != os.path.realpath(desktop):
        try:
            _build_doc().save(desk_path)
        except Exception:
            pass  # Desktop copy is best-effort

    return out_path


# ─────────────────────────────────────────
# PUBLIC ENTRY POINT
# ─────────────────────────────────────────

def solve_assignment(command: str) -> str:
    """
    Main function called from main.py.
    Parses command → searches entire filesystem → reads file → AI completes → saves result.
    """
    # 1. Extract filename from command
    filename = parse_filename(command)
    if not filename:
        return (
            "❌ Couldn't detect a filename.\n"
            "Try: 'complete assignment homework.txt' or "
            "'do my assignment Experiment_8.docx'"
        )

    # 2. Search entire filesystem for the file
    print(f"🔍 Searching for '{filename}' across your PC...")
    filepath = _find_file(filename)

    if not filepath:
        desktop = _get_real_desktop()
        print(f"   Desktop resolved to: {desktop}")
        try:
            files = [f for f in os.listdir(desktop) if os.path.isfile(os.path.join(desktop, f))]
            file_list = ", ".join(files[:10]) if files else "no files found"
        except Exception:
            file_list = "could not list"

        return (
            f"❌ File '{filename}' was not found anywhere on your PC.\n"
            f"📂 Desktop ({desktop}) contains: {file_list}\n"
            f"💡 Make sure the filename is spelled correctly."
        )

    found_name = os.path.basename(filepath)
    print(f"✅ Found: {filepath}")

    # 3. Read the file
    try:
        content = _read_file(filepath)
    except RuntimeError as e:
        return f"❌ Could not read file: {e}"
    except Exception as e:
        return f"❌ Unexpected error reading file: {e}"

    if not content.strip():
        return f"❌ File '{found_name}' appears to be empty — nothing to complete."

    # 4. Ask AI to complete the assignment
    print(f"📄 Read {len(content)} characters from '{found_name}'")
    print(f"🤖 Sending to AI for completion...")

    try:
        result = _ask_ai(content, found_name)
    except Exception as e:
        return f"❌ AI error: {e}"

    # 5. Save result
    try:
        out_path  = _save_result(filepath, found_name, result)
        saved_name = os.path.basename(out_path)
    except Exception as e:
        return (
            f"✅ Assignment completed for '{found_name}'!\n"
            f"⚠️ Could not save file: {e}\n\n"
            f"{'='*50}\n{result}"
        )

    return (
        f"✅ Assignment '{found_name}' completed!\n"
        f"📍 Found at: {filepath}\n"
        f"💾 Saved as: {saved_name} (Word .docx — next to original + copy on Desktop)\n"
        f"{'─'*50}\n"
        f"{result[:900]}{'...[see saved .docx for full answer]' if len(result) > 900 else ''}\n"
        f"{'─'*50}"
    )